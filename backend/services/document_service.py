"""
Document parsing service for PDF, Word, Excel, and JSON files.
"""

import json
from io import BytesIO
from typing import Optional

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

try:
    import pandas as pd
except ImportError:
    pd = None


def parse_document_bytes(
    data: bytes,
    mimetype: str,
    filename: Optional[str] = None,
) -> dict:
    """
    Parse document content from various formats.
    
    Returns:
    {
        "filename": str,
        "mimetype": str,
        "content": str,  # Extracted text content
        "metadata": dict,  # Additional metadata like page count, sheets, etc.
    }
    """
    filename = filename or "document"
    content = ""
    metadata = {}
    
    # Determine document type
    mime_lower = (mimetype or "").lower()
    fname_lower = filename.lower()
    
    try:
        # PDF parsing
        if "pdf" in mime_lower or fname_lower.endswith(".pdf"):
            if PdfReader is None:
                raise RuntimeError("pypdf not installed. Run: pip install pypdf")
            
            pdf = PdfReader(BytesIO(data))
            pages = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i+1}]\n{text.strip()}")
            
            content = "\n\n".join(pages)
            metadata = {
                "page_count": len(pdf.pages),
                "format": "PDF"
            }
        
        # Word document parsing (.docx)
        elif "wordprocessing" in mime_lower or "msword" in mime_lower or fname_lower.endswith((".docx", ".doc")):
            if DocxDocument is None:
                raise RuntimeError("python-docx not installed. Run: pip install python-docx")
            
            doc = DocxDocument(BytesIO(data))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "format": "Word Document"
            }
        
        # Excel parsing (.xlsx, .xls)
        elif "spreadsheet" in mime_lower or fname_lower.endswith((".xlsx", ".xls")):
            if load_workbook is None or pd is None:
                raise RuntimeError("openpyxl and pandas not installed. Run: pip install openpyxl pandas")
            
            wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
            sheets = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_str = "\t".join(str(cell) if cell is not None else "" for cell in row)
                        rows.append(row_str)
                
                if rows:
                    sheet_content = f"[Sheet: {sheet_name}]\n" + "\n".join(rows[:100])  # Limit to 100 rows per sheet
                    sheets.append(sheet_content)
            
            content = "\n\n".join(sheets)
            metadata = {
                "sheet_count": len(wb.sheetnames),
                "sheet_names": wb.sheetnames,
                "format": "Excel Spreadsheet"
            }
        
        # JSON parsing
        elif "json" in mime_lower or fname_lower.endswith(".json"):
            json_data = json.loads(data.decode("utf-8"))
            # Pretty print JSON with indentation
            content = json.dumps(json_data, indent=2, ensure_ascii=False)
            metadata = {
                "format": "JSON",
                "type": type(json_data).__name__
            }
        
        # Plain text fallback
        elif "text" in mime_lower or fname_lower.endswith((".txt", ".csv", ".md")):
            content = data.decode("utf-8", errors="ignore")
            metadata = {
                "format": "Plain Text",
                "encoding": "utf-8"
            }
        
        else:
            raise ValueError(f"Unsupported document type: {mimetype}")
        
        # Truncate if too large (limit to ~100k chars)
        if len(content) > 100_000:
            content = content[:100_000] + "\n\n[Content truncated due to size...]"
            metadata["truncated"] = True
        
        return {
            "filename": filename,
            "mimetype": mimetype,
            "content": content,
            "metadata": metadata,
        }
    
    except Exception as e:
        raise RuntimeError(f"Failed to parse document: {str(e)}")
